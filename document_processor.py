
"""
Document Processor Module
=========================

This module handles extracting text from different document types:
- PDF files (with OCR support for image-based PDFs)
- DOCX files (Microsoft Word documents)
- TXT files (plain text files)

Author: Your Name
Date: June 2025
"""

import logging
import re
from typing import Optional, Dict, Any
import pymupdf  # For PDF processing
import pytesseract  # For OCR (Optical Character Recognition)
from PIL import Image  # For image processing
import io
from docx import Document  # For DOCX processing

# Configure logging to help with debugging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    A class to handle text extraction from various document formats.

    This class provides methods to extract text from PDF, DOCX, and TXT files.
    It also includes OCR capabilities for image-based PDFs.
    """

    def __init__(self):
        """Initialize the DocumentProcessor."""
        logger.info("DocumentProcessor initialized")

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from a PDF file.

        This method first tries to extract text directly from the PDF.
        If no text is found (image-based PDF), it uses OCR to extract text.

        Args:
            file_content (bytes): The PDF file content as bytes

        Returns:
            str: Extracted text from the PDF
        """
        try:
            logger.info("Starting PDF text extraction")

            # Open the PDF document from bytes
            doc = pymupdf.open(stream=file_content, filetype="pdf")
            text = ""

            # First, try to extract text directly from each page
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                text += page_text + "\n"

            # If we got very little text, the PDF might be image-based
            # Let's try OCR on each page
            if len(text.strip()) < 50:  # If less than 50 characters
                logger.info("Little text found, trying OCR")
                text = ""

                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    # Convert page to image
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")

                    # Use OCR to extract text from the image
                    try:
                        img = Image.open(io.BytesIO(img_data))
                        ocr_text = pytesseract.image_to_string(img)
                        text += ocr_text + "\n"
                    except Exception as ocr_error:
                        logger.warning(f"OCR failed for page {page_num}: {str(ocr_error)}")
                        continue

            doc.close()
            logger.info(f"PDF text extraction completed. Extracted {len(text)} characters")
            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise Exception(f"Failed to extract text from PDF: {str(e)}")

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from a DOCX file.

        Args:
            file_content (bytes): The DOCX file content as bytes

        Returns:
            str: Extracted text from the DOCX file
        """
        try:
            logger.info("Starting DOCX text extraction")

            # Open the DOCX document from bytes
            doc = Document(io.BytesIO(file_content))
            text = ""

            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Also extract text from tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"

            logger.info(f"DOCX text extraction completed. Extracted {len(text)} characters")
            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")

    def extract_text_from_txt(self, file_content: bytes) -> str:
        """
        Extract text from a TXT file.

        Args:
            file_content (bytes): The TXT file content as bytes

        Returns:
            str: The text content of the file
        """
        try:
            logger.info("Starting TXT text extraction")

            # Try different encodings to handle various text files
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    logger.info(f"TXT text extraction completed using {encoding}. Extracted {len(text)} characters")
                    return text.strip()
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, try with error handling
            text = file_content.decode('utf-8', errors='ignore')
            logger.warning("Used UTF-8 with error ignoring for TXT extraction")
            return text.strip()

        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            raise Exception(f"Failed to extract text from TXT: {str(e)}")

    def process_document(self, file_content: bytes, file_type: str) -> str:
        """
        Process a document and extract text based on its type.

        Args:
            file_content (bytes): The file content as bytes
            file_type (str): The type of file ('pdf', 'docx', or 'txt')

        Returns:
            str: Extracted text from the document
        """
        logger.info(f"Processing document of type: {file_type}")

        if file_type.lower() == 'pdf':
            return self.extract_text_from_pdf(file_content)
        elif file_type.lower() == 'docx':
            return self.extract_text_from_docx(file_content)
        elif file_type.lower() == 'txt':
            return self.extract_text_from_txt(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

# Test the module if run directly
if __name__ == "__main__":
    processor = DocumentProcessor()
    print("DocumentProcessor module loaded successfully!")